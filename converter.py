from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
import requests


def html_to_docx(request):
    # Retrieve the HTML content
    url = "https://www.freecodecamp.org/news/html-tables-table-tutorial-with-css-example-code/"
    response = requests.get(url)
    html_content = response.content

    # Convert the HTML content to a DOCX file
    document = Document()
    soup = BeautifulSoup(html_content, "html.parser")
    for tag in soup.find_all():
        if tag.name == "img":
            image_content = requests.get(tag["src"]).content
            with open("image.jpg", "wb") as f:
                f.write(image_content)
            document.add_picture("image.jpg", width=Inches(6))
        elif tag.name == "p":
            document.add_paragraph(tag.text)
    document.save("example.docx")

    # Return the DOCX file as a response
    with open("example.docx", "rb") as f:
        response = HttpResponse(
            f.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = "attachment; filename=example.docx"
        return response


html_to_docx(requests)
